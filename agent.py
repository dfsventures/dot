import os, json, base64, pickle, io, asyncio, glob, re, shutil
from anthropic import Anthropic
from telegram import Update
from telegram.error import BadRequest
from telegram.ext import ApplicationBuilder, MessageHandler, CommandHandler, filters, ContextTypes
from dotenv import load_dotenv
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import dropbox as dbx_lib
import requests
import logging

logging.basicConfig(level=logging.INFO)

load_dotenv('.dot.env')

TELEGRAM_TOKEN    = os.getenv("TELEGRAM_TOKEN")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
YOUR_USER_ID      = int(os.getenv("YOUR_TELEGRAM_USER_ID"))
GRANOLA_TOKEN     = os.getenv("GRANOLA_TOKEN")
DROPBOX_TOKEN     = os.getenv("DROPBOX_TOKEN")

client = Anthropic(api_key=ANTHROPIC_API_KEY)

# ── GOOGLE AUTH ───────────────────────────────────────────────────────────────
def load_google_creds(pickle_path):
    with open(pickle_path, 'rb') as f:
        creds = pickle.load(f)
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
        with open(pickle_path, 'wb') as f:
            pickle.dump(creds, f)
    return creds

work_creds     = load_google_creds('token_work.pickle')
gmail_work     = build('gmail',    'v1', credentials=work_creds)
calendar_work  = build('calendar', 'v3', credentials=work_creds)
drive_work     = build('drive',    'v3', credentials=work_creds)

# ── DROPBOX ───────────────────────────────────────────────────────────────────
dbx = dbx_lib.Dropbox(
    oauth2_access_token=os.getenv("DROPBOX_TOKEN"),
    oauth2_refresh_token=os.getenv("DROPBOX_REFRESH_TOKEN"),
    app_key=os.getenv("DROPBOX_APP_KEY"),
    app_secret=os.getenv("DROPBOX_APP_SECRET")
)
# ── MEMORY (shared with ingest.py) ────────────────────────────────────────────
from memory import (
    save_memory, delete_memory, get_all_memories,
    retrieve_relevant_memories, migrate_sqlite_to_chroma, parse_json_array, response_text,
    response_text_checked,
    upsert_deal as _upsert_deal, get_deal as _get_deal, list_deals as _list_deals,
    save_procedure as _save_procedure, get_all_procedures, delete_procedure as _delete_procedure,
    retrieve_relevant_procedures, get_cached_doc, save_cached_doc,
    was_prepped, mark_prepped, add_prep_mute, list_prep_mutes, delete_prep_mute, is_prep_muted,
    save_feedback, search_bulk_records,
)

# Run migration on startup to catch any memories added before vector search
migrate_sqlite_to_chroma()

async def send_markdown(sender, text: str):
    """sender: async callable shaped like `reply_text`/`send_message` — takes
    (text, parse_mode=...). Tries Telegram Markdown so the model's **bold**/bullets
    actually render instead of showing as literal asterisks; falls back to plain
    text if the output isn't valid Markdown (an unmatched */_/`/[ would otherwise
    make Telegram reject the whole message)."""
    try:
        await sender(text, parse_mode="Markdown")
    except BadRequest:
        await sender(text, parse_mode=None)

def extract_and_save_memories(conversation):
    if len(conversation) < 4:
        return
    convo_text = "\n".join([
        f"{m['role'].upper()}: {m['content'] if isinstance(m['content'], str) else '[tool call]'}"
        for m in conversation
    ])
    try:
        response = client.messages.create(
            model="claude-sonnet-5", max_tokens=1500,
            thinking={"type": "disabled"},
            system='Extract facts worth remembering. Return ONLY a JSON array of strings. If nothing, return []. Focus on people, companies, deals, preferences.',
            messages=[{"role": "user", "content": convo_text}]
        )
        facts = parse_json_array(response_text_checked(response, "extract_and_save_memories"))
        print(f"Extracted facts: {facts}")
        for fact in facts:
            if fact and len(fact) > 10:
                save_memory(fact)
    except Exception as e:
        print(f"Memory error: {e}")

# ── CONTEXT WINDOW MANAGEMENT ─────────────────────────────────────────────────
# Sonnet 5 has a 1M token context window — we manage at 80k tokens to keep
# costs reasonable and latency low for a chat agent use case.
CONTEXT_TOKEN_LIMIT = 80_000

def estimate_tokens(messages: list) -> int:
    """Rough token estimate: ~1 token per 4 chars of text content."""
    total = 0
    for m in messages:
        content = m.get("content", "")
        if isinstance(content, str):
            total += len(content) // 4
        elif isinstance(content, list):
            for block in content:
                if isinstance(block, dict):
                    total += len(str(block)) // 4
    return total

def compress_history(conversation: list) -> list:
    """Summarise all but the last few messages into a single context message.

    The kept window must start at a plain-text user message — cutting at a
    tool_result message would orphan it from its assistant tool_use turn and
    the API would reject the request.
    """
    if len(conversation) <= 4:
        return conversation
    cut = len(conversation) - 4
    while cut > 0 and not (
        conversation[cut].get("role") == "user"
        and isinstance(conversation[cut].get("content"), str)
    ):
        cut -= 1
    if cut == 0:
        return conversation  # no safe boundary found — skip this round
    to_summarise = conversation[:cut]
    recent = conversation[cut:]
    convo_text = "\n".join([
        f"{m['role'].upper()}: {m['content'] if isinstance(m['content'], str) else '[tool interaction]'}"
        for m in to_summarise
    ])
    try:
        r = client.messages.create(
            model="claude-sonnet-5", max_tokens=1500,
            thinking={"type": "disabled"},
            system="Summarise this conversation history concisely, preserving all key facts, decisions, and context that would be needed to continue the conversation intelligently. Be dense — this replaces the full history.",
            messages=[{"role": "user", "content": convo_text}]
        )
        summary = response_text_checked(r, "compress_history").strip()
        if not summary:
            return conversation  # no summary → keep everything, never silently truncate
        summary_message = {
            "role": "user",
            "content": f"[Earlier conversation summary]\n{summary}"
        }
        ack = {
            "role": "assistant",
            "content": "Understood, I have the earlier context."
        }
        logging.info(f"Context compressed: {len(to_summarise)} messages → summary")
        return [summary_message, ack] + recent
    except Exception as e:
        logging.error(f"Compression error: {e}")
        return conversation  # was: return recent — dropped all but 4 messages on any failure

# ── GMAIL HELPERS ─────────────────────────────────────────────────────────────
def _search_gmail_svc(svc, query, max_results=10):
    results = svc.users().messages().list(userId='me', q=query, maxResults=max_results).execute()
    messages = results.get('messages', [])
    if not messages:
        return "No emails found."
    output = []
    for msg in messages[:max_results]:
        m = svc.users().messages().get(userId='me', id=msg['id'], format='metadata',
            metadataHeaders=['Subject', 'From', 'Date']).execute()
        headers = {h['name']: h['value'] for h in m['payload']['headers']}
        output.append(f"From: {headers.get('From','?')} | Date: {headers.get('Date','?')} | Subject: {headers.get('Subject','?')} | ID: {msg['id']}")
    return "\n".join(output)

def _walk_payload_for_text(payload, mime="text/plain"):
    """Depth-first search of a Gmail payload tree for the first body of the given mime type."""
    if payload.get('mimeType') == mime and payload.get('body', {}).get('data'):
        return base64.urlsafe_b64decode(payload['body']['data']).decode('utf-8', errors='ignore')
    for part in payload.get('parts', []):
        text = _walk_payload_for_text(part, mime)
        if text:
            return text
    return ""

def _read_gmail_svc(svc, message_id):
    m = svc.users().messages().get(userId='me', id=message_id, format='full').execute()
    body = _walk_payload_for_text(m['payload'])
    if not body:
        html = _walk_payload_for_text(m['payload'], mime="text/html")
        if html:
            body = re.sub(r'<[^>]+>', ' ', html)  # crude tag strip — good enough for reading
    return body[:3000] if body.strip() else "[No text body found]"

# ── TOOL FUNCTIONS ────────────────────────────────────────────────────────────
def search_gmail_work(query: str, max_results: int = 10):
    return _search_gmail_svc(gmail_work, query, max_results)

def read_gmail_work(message_id: str):
    return _read_gmail_svc(gmail_work, message_id)

def create_gmail_draft(to: str, subject: str, body: str, thread_id: str = None) -> str:
    from email.mime.text import MIMEText
    msg = MIMEText(body)
    msg['to'] = to
    msg['subject'] = subject
    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
    draft_body = {'message': {'raw': raw}}
    if thread_id:
        draft_body['message']['threadId'] = thread_id
    try:
        result = gmail_work.users().drafts().create(userId='me', body=draft_body).execute()
        return f"Draft created (ID: {result['id']}). Open Gmail to review and send."
    except Exception as e:
        return f"Draft creation error: {e}"

def list_calendar_events(days_ahead: int = 7, max_results: int = 10):
    """List upcoming calendar events from work Google Calendar."""
    from datetime import datetime, timezone, timedelta
    now = datetime.now(timezone.utc)
    end = now + timedelta(days=days_ahead)
    try:
        events_result = calendar_work.events().list(
            calendarId='primary',
            timeMin=now.isoformat(),
            timeMax=end.isoformat(),
            maxResults=max_results,
            singleEvents=True,
            orderBy='startTime'
        ).execute()
        events = events_result.get('items', [])
        if not events:
            return f"No events in the next {days_ahead} days."
        lines = []
        for e in events:
            start = e['start'].get('dateTime', e['start'].get('date', ''))[:16].replace('T', ' ')
            title = e.get('summary', 'Untitled')
            attendees = e.get('attendees', [])
            attendee_str = ""
            if attendees:
                names = [a.get('displayName') or a.get('email', '') for a in attendees[:4]]
                attendee_str = f" | With: {', '.join(names)}"
            event_id = e.get('id', '')
            lines.append(f"{start} — {title}{attendee_str} | ID: {event_id}")
        return "\n".join(lines)
    except Exception as e:
        return f"Calendar error: {e}"

def get_calendar_event(event_id: str):
    """Get full details of a specific calendar event including description and all attendees."""
    try:
        e = calendar_work.events().get(calendarId='primary', eventId=event_id).execute()
        lines = []
        lines.append(f"Title: {e.get('summary', 'Untitled')}")
        start = e['start'].get('dateTime', e['start'].get('date', ''))
        lines.append(f"Start: {start}")
        end = e['end'].get('dateTime', e['end'].get('date', ''))
        lines.append(f"End: {end}")
        location = e.get('location', '')
        if location:
            lines.append(f"Location: {location}")
        attendees = e.get('attendees', [])
        if attendees:
            att_list = [f"{a.get('displayName', '')} <{a.get('email', '')}> ({'accepted' if a.get('responseStatus') == 'accepted' else a.get('responseStatus', '?')})" for a in attendees]
            lines.append(f"Attendees:\n  " + "\n  ".join(att_list))
        desc = e.get('description', '')
        if desc:
            lines.append(f"Description:\n{desc[:500]}")
        meet_link = e.get('hangoutLink', '') or (e.get('conferenceData') or {}).get('entryPoints', [{}])[0].get('uri', '')
        if meet_link:
            lines.append(f"Meeting link: {meet_link}")
        return "\n".join(lines)
    except Exception as e:
        return f"Calendar event error: {e}"

def create_calendar_event(
    title: str,
    start: str,
    end: str,
    attendees: list = None,
    description: str = "",
    location: str = "",
    add_meet_link: bool = False
):
    """Create a new event on Joey's work Google Calendar.
    start/end must be ISO 8601 format: '2026-06-15T14:00:00+01:00' (with timezone offset).
    attendees is a list of email address strings.
    """
    from datetime import datetime, timezone
    try:
        event = {
            "summary": title,
            "start":   {"dateTime": start, "timeZone": "America/Toronto"},
            "end":     {"dateTime": end,   "timeZone": "America/Toronto"},
        }
        if description:
            event["description"] = description
        if location:
            event["location"] = location
        if attendees:
            event["attendees"] = [{"email": a} for a in attendees]
        if add_meet_link:
            import uuid
            event["conferenceData"] = {
                "createRequest": {
                    "requestId": str(uuid.uuid4()),
                    "conferenceSolutionKey": {"type": "hangoutsMeet"}
                }
            }
        kwargs = {"calendarId": "primary", "body": event, "sendUpdates": "all"}
        if add_meet_link:
            kwargs["conferenceDataVersion"] = 1
        created = calendar_work.events().insert(**kwargs).execute()
        link = created.get("hangoutLink") or created.get("htmlLink", "")
        return f"Event created: {created.get('summary')} on {created['start'].get('dateTime', '')[:16]} | Link: {link} | ID: {created['id']}"
    except Exception as e:
        return f"Calendar create error: {e}"

def delete_calendar_event(event_id: str):
    """Delete a calendar event by ID. Use get_calendar_event first to confirm the right event."""
    try:
        calendar_work.events().delete(calendarId='primary', eventId=event_id, sendUpdates='all').execute()
        return f"Event {event_id} deleted and attendees notified."
    except Exception as e:
        return f"Calendar delete error: {e}"

def update_calendar_event(
    event_id: str,
    title: str = None,
    start: str = None,
    end: str = None,
    description: str = None,
    location: str = None,
    add_attendees: list = None
):
    """Update fields on an existing calendar event. Only provided fields are changed."""
    try:
        event = calendar_work.events().get(calendarId='primary', eventId=event_id).execute()
        if title:
            event['summary'] = title
        if start:
            event['start'] = {"dateTime": start, "timeZone": "America/Toronto"}
        if end:
            event['end'] = {"dateTime": end, "timeZone": "America/Toronto"}
        if description is not None:
            event['description'] = description
        if location is not None:
            event['location'] = location
        if add_attendees:
            existing = event.get('attendees', [])
            existing_emails = {a['email'] for a in existing}
            for email in add_attendees:
                if email not in existing_emails:
                    existing.append({"email": email})
            event['attendees'] = existing
        updated = calendar_work.events().update(
            calendarId='primary', eventId=event_id, body=event, sendUpdates='all'
        ).execute()
        return f"Event updated: {updated.get('summary')} | {updated['start'].get('dateTime','')[:16]}"
    except Exception as e:
        return f"Calendar update error: {e}"

def search_granola(query: str, max_results: int = 10):
    headers = {"Authorization": f"Bearer {GRANOLA_TOKEN}"}
    try:
        r = requests.get(
            "https://public-api.granola.ai/v1/notes",
            headers=headers,
            params={"page_size": 30},
            timeout=10
        )
        if r.status_code == 401:
            return "Granola auth error: check your API token and that your plan is Enterprise."
        if r.status_code != 200:
            return f"Granola API error: {r.status_code} {r.text[:200]}"
        notes = r.json().get("notes", [])
        if not notes:
            return "No Granola notes found."
        query_lower = query.lower()
        def _note_matches(n):
            if query_lower in (n.get("title") or "").lower():
                return True
            for a in n.get("attendees", []):
                if query_lower in (a.get("name") or "").lower():
                    return True
                if query_lower in (a.get("email") or "").lower():
                    return True
            return False
        matched = [n for n in notes if _note_matches(n)]
        results = matched if matched else notes
        lines = []
        for n in results[:max_results]:
            title = n.get("title") or "Untitled"
            date  = (n.get("created_at") or "")[:10]
            note_id = n.get("id", "")
            lines.append(f"Title: {title} | Date: {date} | ID: {note_id}")
        return "\n".join(lines)
    except Exception as e:
        return f"Granola search error: {e}"

def read_granola(note_id: str):
    headers = {"Authorization": f"Bearer {GRANOLA_TOKEN}"}
    try:
        r = requests.get(
            f"https://public-api.granola.ai/v1/notes/{note_id}",
            headers=headers,
            params={"include": "transcript"},
            timeout=10
        )
        if r.status_code == 404:
            return "Note not found — it may not have a generated summary yet."
        if r.status_code != 200:
            return f"Granola read error: {r.status_code}"
        data = r.json()
        output = []
        attendees = data.get("attendees", [])
        if attendees:
            names = ", ".join(a.get("name") or a.get("email", "") for a in attendees)
            output.append(f"Attendees: {names}")
        start = ((data.get("calendar_event") or {}).get("scheduled_start_time") or "")[:10]
        if start:
            output.append(f"Date: {start}")
        summary = data.get("summary_text") or data.get("summary_markdown") or ""
        if summary:
            output.append(f"\nSummary:\n{summary[:2000]}")
        transcript = data.get("transcript") or []
        if transcript:
            lines = []
            for t in transcript[:40]:
                source = (t.get("speaker") or {}).get("source", "?")
                speaker = "You" if source == "microphone" else "Them"
                lines.append(f"{speaker}: {t.get('text','')}")
            output.append(f"\nTranscript (first 40 turns):\n" + "\n".join(lines))
        return "\n".join(output)[:3000] if output else "[Note has no content]"
    except Exception as e:
        return f"Granola read error: {e}"

def search_dropbox(query: str, max_results: int = 10):
    import dropbox.files as dbx_files
    try:
        options = dbx_files.SearchOptions(max_results=max_results, filename_only=True)
        results = dbx.files_search_v2(query, options=options)
        lines = []
        for m in results.matches:
            try:
                wrapper = m.metadata
                # MetadataV2 is a union type — unwrap only if the tag is "metadata"
                if hasattr(wrapper, 'is_metadata') and wrapper.is_metadata():
                    meta = wrapper.get_metadata()
                else:
                    meta = wrapper
                # Guard against SDK versions that return a bound method instead of the value
                if callable(meta) or not hasattr(meta, 'name'):
                    continue
                size_kb = getattr(meta, 'size', 0) // 1024
                lines.append(f"Name: {meta.name} | Path: {meta.path_display} | Size: {size_kb}KB")
            except Exception:
                continue
        return "\n".join(lines) if lines else "No Dropbox files found."
    except Exception as e:
        return f"Dropbox search error: {e}"

def update_deal(company: str, stage: str = None, last_touchpoint: str = None,
               next_action: str = None, notes: str = None) -> str:
    deal = _upsert_deal(company, stage, last_touchpoint, next_action, notes)
    return (f"Deal: {deal['company']} | Stage: {deal['stage']} | "
            f"Next action: {deal['next_action'] or '—'} | "
            f"Last touchpoint: {deal['last_touchpoint'] or '—'}")

def get_deal_info(company: str) -> str:
    deal = _get_deal(company)
    if not deal:
        return f"No deal found for '{company}'."
    lines = [
        f"Company: {deal['company']}",
        f"Stage: {deal['stage']}",
        f"Last touchpoint: {deal['last_touchpoint'] or '—'}",
        f"Next action: {deal['next_action'] or '—'}",
        f"Notes: {deal['notes'] or '—'}",
        f"Created: {deal['created_at'][:10]} | Updated: {deal['updated_at'][:10]}",
    ]
    from memory import conn as _mem_conn
    rows = _mem_conn.execute(
        "SELECT content FROM memories WHERE tags LIKE ? ORDER BY id DESC LIMIT 10",
        (f"%deal:{deal['company']}%",),
    ).fetchall()
    if rows:
        lines.append("\nFrom ingested documents:")
        lines.extend(f"- {r[0]}" for r in rows)
    return "\n".join(lines)

def list_deals(stage: str = None) -> str:
    deals = _list_deals(stage)
    if not deals:
        return f"No deals{' in stage: ' + stage if stage else ' in pipeline'}."
    return "\n".join(
        f"{d['company']} | {d['stage']} | Next: {d['next_action'] or '—'} | Updated: {d['updated_at'][:10]}"
        for d in deals
    )

def search_memory(query: str) -> str:
    results = retrieve_relevant_memories(query, k=30)
    if not results:
        return "No memories found matching that query."
    return "\n".join(f"• {m}" for m in results)

def search_deal_database(query: str, max_results: int = 10) -> str:
    results = search_bulk_records(query, k=max_results)
    if not results:
        return "No matching rows found in the deal database."
    return "\n".join(f"• {r}" for r in results)

def save_procedure(trigger: str, procedure: str) -> str:
    _save_procedure(trigger, procedure)
    return f"Saved procedure for future '{trigger}' situations."

def mute_meeting_prep(pattern: str, reason: str = "") -> str:
    """Stop sending automatic prep briefs for a meeting. `pattern` is an event title or a
    distinctive substring of one. Never raises — returns a string like every other tool."""
    try:
        add_prep_mute(pattern, reason)
        return f"Muted automatic prep briefs for anything matching '{pattern}'. This is persistent."
    except Exception as e:
        return f"Mute error: {e}"

def list_meeting_prep_mutes() -> str:
    try:
        mutes = list_prep_mutes()
        if not mutes:
            return "No meeting prep mutes set."
        return "\n".join(
            f"{m['id']}. '{m['pattern']}'" + (f" — {m['reason']}" if m['reason'] else "")
            for m in mutes
        )
    except Exception as e:
        return f"List mutes error: {e}"

def unmute_meeting_prep(mute_id: int) -> str:
    try:
        ok = delete_prep_mute(int(mute_id))
        return f"Unmuted #{mute_id}." if ok else f"No mute found with id {mute_id}."
    except Exception as e:
        return f"Unmute error: {e}"

_MIN_TEXT_CHARS = 200  # same threshold ingest.py uses to detect an image-only PDF

def _pdf_text_or_marker(content: bytes, label: str) -> str:
    """Extract a PDF's text layer, or return an explicit marker for image-only PDFs."""
    import PyPDF2
    reader = PyPDF2.PdfReader(io.BytesIO(content))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if len(text.strip()) < _MIN_TEXT_CHARS:
        return (f"[{label}: {len(reader.pages)}-page PDF with no text layer (image-based deck). "
                f"Nothing could be extracted locally. Check long-term memory for facts already "
                f"ingested from this file before telling Joey anything about its contents.]")
    return text

_VISION_TRANSCRIPTION_NOTE = (
    "[Transcribed from an image-based PDF by Claude — text is a transcription, "
    "not the original text layer.]\n\n"
)

def _vision_fallback_or_marker(content: bytes, marker: str, source: str, file_key: str,
                               revision: str, filename: str) -> str:
    """WS-13: on a WS-10 'no text layer' marker, try a one-shot live vision transcription of
    an image-only PDF that was never ingested — gated by ingest.py's own size/page caps.
    Caches the raw transcript (kind='vision') so a second read never re-calls Claude; the
    disclaimer note is prepended only to this immediate return, not to the cached copy, so a
    later cache hit stays reusable byte-for-byte. Falls back to the marker, uncached, if the
    transcription is empty (oversized, too many pages, or an API failure)."""
    from ingest import transcribe_pdf_with_claude
    vision_text = transcribe_pdf_with_claude(content, filename)
    if not vision_text:
        return marker
    save_cached_doc(source, file_key, revision, filename, vision_text, kind="vision")
    return _VISION_TRANSCRIPTION_NOTE + vision_text

def read_dropbox_file(file_path: str, offset: int = 0):
    try:
        meta = dbx.files_get_metadata(file_path)   # ~0.2-0.3s, no bytes transferred
        cached = get_cached_doc("dropbox", meta.id, meta.content_hash)
        if cached:
            text = cached
        else:
            metadata, response = dbx.files_download(file_path)
            content = response.content
            name = metadata.name.lower()
            if name.endswith(('.txt', '.md')):
                text = content.decode('utf-8', errors='ignore')
            elif name.endswith('.pdf'):
                try:
                    text = _pdf_text_or_marker(content, metadata.name)
                except Exception as e:
                    return f"[PDF read error: {e}]"
                if text.startswith("["):
                    text = _vision_fallback_or_marker(content, text, "dropbox", meta.id,
                                                       meta.content_hash, metadata.name)
            elif name.endswith('.docx'):
                from docx import Document
                doc = Document(io.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs)
            else:
                return f"[File type not readable as text: {name}]"
            save_cached_doc("dropbox", meta.id, meta.content_hash, metadata.name, text)
        chunk = text[offset:offset + 3000]
        if offset + 3000 < len(text):
            chunk += f"\n[... {len(text) - offset - 3000} more characters — call again with offset={offset + 3000}]"
        return chunk
    except Exception as e:
        return f"Dropbox read error: {e}"

def search_drive(query: str, max_results: int = 10):
    try:
        safe_q = query.replace("'", "\\'")
        results = drive_work.files().list(
            q=f"(name contains '{safe_q}' or fullText contains '{safe_q}') and trashed = false",
            pageSize=max_results, orderBy="modifiedTime desc",
            fields="files(id, name, mimeType, modifiedTime)"
        ).execute()
        files = results.get('files', [])
        if not files:
            return "No Drive files found."
        return "\n".join(
            f"Name: {f['name']} | Type: {f['mimeType'].split('.')[-1]} | "
            f"Modified: {f.get('modifiedTime', '')[:10]} | ID: {f['id']}"
            for f in files
        )
    except Exception as e:
        return f"Drive search error: {e}"

_DRIVE_EXPORTS = {
    'application/vnd.google-apps.document':     'text/plain',
    'application/vnd.google-apps.spreadsheet':  'text/csv',
    'application/vnd.google-apps.presentation': 'text/plain',
}

def read_drive_file(file_id: str, offset: int = 0):
    try:
        meta = drive_work.files().get(fileId=file_id, fields="name, mimeType, version, modifiedTime").execute()
        name, mime = meta['name'].lower(), meta['mimeType']
        cached = get_cached_doc("drive", file_id, meta['version'])
        if cached:
            text = cached
        else:
            if mime in _DRIVE_EXPORTS:
                content = drive_work.files().export(fileId=file_id, mimeType=_DRIVE_EXPORTS[mime]).execute()
                text = content.decode('utf-8', errors='ignore') if isinstance(content, bytes) else str(content)
            else:
                content = drive_work.files().get_media(fileId=file_id).execute()
                if name.endswith(('.txt', '.md', '.csv')):
                    text = content.decode('utf-8', errors='ignore')
                elif name.endswith('.pdf'):
                    text = _pdf_text_or_marker(content, meta['name'])
                    if text.startswith("["):
                        text = _vision_fallback_or_marker(content, text, "drive", file_id,
                                                           meta['version'], meta['name'])
                elif name.endswith('.docx'):
                    from docx import Document
                    doc = Document(io.BytesIO(content))
                    text = "\n".join(p.text for p in doc.paragraphs)
                else:
                    return f"[File type not readable as text: {meta['name']}]"
            save_cached_doc("drive", file_id, meta['version'], meta['name'], text)
        chunk = text[offset:offset + 3000]
        if offset + 3000 < len(text):
            chunk += f"\n[... {len(text) - offset - 3000} more characters — call again with offset={offset + 3000}]"
        return chunk
    except Exception as e:
        return f"Drive read error: {e}"

# ── TOOLS SCHEMA ──────────────────────────────────────────────────────────────
TOOLS = [
    {"type": "web_search_20260209", "name": "web_search"},
    {
        "name": "search_gmail_work",
        "description": "Search Joey's work Google Workspace inbox. Use for founder emails, investor correspondence, DFS Lab work. Gmail search syntax: from:, subject:, after:YYYY/MM/DD.",
        "input_schema": {"type": "object", "properties": {
            "query": {"type": "string"},
            "max_results": {"type": "integer", "default": 10}
        }, "required": ["query"]}
    },
    {
        "name": "read_gmail_work",
        "description": "Read a work Gmail message body by ID (get ID from search_gmail_work first).",
        "input_schema": {"type": "object", "properties": {"message_id": {"type": "string"}}, "required": ["message_id"]}
    },
    {
        "name": "create_gmail_draft",
        "description": "Create a Gmail draft for Joey to review and send himself — never sends automatically. Use when asked to draft or reply to an email. Read the thread first with search_gmail_work + read_gmail_work so the draft has full context. For replies, pass the thread_id so the draft appears in the right thread.",
        "input_schema": {"type": "object", "properties": {
            "to":        {"type": "string", "description": "Recipient email address"},
            "subject":   {"type": "string", "description": "Subject line — use 'Re: original subject' for replies"},
            "body":      {"type": "string", "description": "Full email body text"},
            "thread_id": {"type": "string", "description": "Gmail thread ID to reply in-thread (get from search_gmail_work results)"}
        }, "required": ["to", "subject", "body"]}
    },
    {
        "name": "list_calendar_events",
        "description": "List upcoming events from Joey's work Google Calendar. Use for 'what's on my calendar', 'what do I have today/this week', 'who am I meeting'. Defaults to next 7 days.",
        "input_schema": {"type": "object", "properties": {
            "days_ahead": {"type": "integer", "default": 7, "description": "How many days ahead to look"},
            "max_results": {"type": "integer", "default": 10}
        }}
    },
    {
        "name": "get_calendar_event",
        "description": "Get full details of a specific calendar event by ID (get ID from list_calendar_events first). Returns all attendees, description, and meeting link.",
        "input_schema": {"type": "object", "properties": {"event_id": {"type": "string"}}, "required": ["event_id"]}
    },
    {
        "name": "create_calendar_event",
        "description": "Create a new event on Joey's work Google Calendar. Use for 'schedule a call', 'block time', 'add a meeting'. Always confirm title, start, and end before calling. start/end format: '2026-06-15T14:00:00-04:00' (include Toronto UTC offset: EDT is -04:00, EST is -05:00).",
        "input_schema": {"type": "object", "properties": {
            "title":         {"type": "string", "description": "Event title"},
            "start":         {"type": "string", "description": "ISO 8601 start datetime with timezone offset"},
            "end":           {"type": "string", "description": "ISO 8601 end datetime with timezone offset"},
            "attendees":     {"type": "array", "items": {"type": "string"}, "description": "List of attendee email addresses"},
            "description":   {"type": "string", "description": "Event description or agenda"},
            "location":      {"type": "string", "description": "Location or video link"},
            "add_meet_link": {"type": "boolean", "description": "If true, auto-generate a Google Meet link", "default": False}
        }, "required": ["title", "start", "end"]}
    },
    {
        "name": "update_calendar_event",
        "description": "Update an existing calendar event. Only fields you provide are changed. Get event_id from list_calendar_events or get_calendar_event first.",
        "input_schema": {"type": "object", "properties": {
            "event_id":      {"type": "string"},
            "title":         {"type": "string"},
            "start":         {"type": "string", "description": "ISO 8601 datetime with timezone offset"},
            "end":           {"type": "string", "description": "ISO 8601 datetime with timezone offset"},
            "description":   {"type": "string"},
            "location":      {"type": "string"},
            "add_attendees": {"type": "array", "items": {"type": "string"}, "description": "Additional attendee emails to add"}
        }, "required": ["event_id"]}
    },
    {
        "name": "delete_calendar_event",
        "description": "Delete a calendar event and notify attendees. Use get_calendar_event to confirm the right event before deleting.",
        "input_schema": {"type": "object", "properties": {"event_id": {"type": "string"}}, "required": ["event_id"]}
    },
    {
        "name": "search_granola",
        "description": "Search Granola meeting notes by title keyword. Returns recent notes matching the query. Use when asked about calls, meetings, or what was discussed with a founder or contact.",
        "input_schema": {"type": "object", "properties": {
            "query": {"type": "string"},
            "max_results": {"type": "integer", "default": 10}
        }, "required": ["query"]}
    },
    {
        "name": "read_granola",
        "description": "Read a full Granola meeting note by ID (get ID from search_granola first). Returns attendees, summary, and transcript.",
        "input_schema": {"type": "object", "properties": {"note_id": {"type": "string"}}, "required": ["note_id"]}
    },
    {
        "name": "search_dropbox",
        "description": "Search Dropbox files by name. Use for pitch decks, PDFs, and documents stored in Dropbox.",
        "input_schema": {"type": "object", "properties": {
            "query": {"type": "string"},
            "max_results": {"type": "integer", "default": 10}
        }, "required": ["query"]}
    },
    {
        "name": "read_dropbox_file",
        "description": "Download and read a Dropbox file by its full path (get path from search_dropbox first). Supports PDF, DOCX, TXT, MD (plus PPTX when previously ingested). Image-only PDFs (scanned/designed decks with no text layer) are transcribed live by Claude on first read and cached — this takes 20-60s the first time, then is instant. Returns up to 3,000 characters starting at offset; if the result says more characters remain, call again with the suggested offset to page further into the document.",
        "input_schema": {"type": "object", "properties": {
            "file_path": {"type": "string"},
            "offset": {"type": "integer", "default": 0}
        }, "required": ["file_path"]}
    },
    {
        "name": "search_drive",
        "description": "Search Joey's Google Drive by file name. Use for documents and shared files that live in Drive rather than Dropbox.",
        "input_schema": {"type": "object", "properties": {
            "query": {"type": "string"},
            "max_results": {"type": "integer", "default": 10}
        }, "required": ["query"]}
    },
    {
        "name": "read_drive_file",
        "description": "Read a Google Drive file by ID (get ID from search_drive first). Supports Google Docs/Sheets/Slides (exported as text) plus PDF, DOCX, TXT, MD, CSV. Image-only PDFs (scanned/designed decks with no text layer) are transcribed live by Claude on first read and cached — this takes 20-60s the first time, then is instant. Returns up to 3,000 characters starting at offset; if the result says more characters remain, call again with the suggested offset to page further into the document.",
        "input_schema": {"type": "object", "properties": {
            "file_id": {"type": "string"},
            "offset": {"type": "integer", "default": 0}
        }, "required": ["file_id"]}
    },
    {
        "name": "update_deal",
        "description": "Create or update a deal in the pipeline. Use for 'add company X', 'move X to due diligence', 'log that I spoke with X', 'set next action for X'. Stages: sourcing, first_call, due_diligence, passed, invested.",
        "input_schema": {"type": "object", "properties": {
            "company":         {"type": "string", "description": "Company name"},
            "stage":           {"type": "string", "description": "sourcing | first_call | due_diligence | passed | invested"},
            "last_touchpoint": {"type": "string", "description": "Brief note on the last interaction, e.g. 'Call 2026-06-11, discussed Series A terms'"},
            "next_action":     {"type": "string", "description": "What needs to happen next, e.g. 'Send term sheet by Jun 20'"},
            "notes":           {"type": "string", "description": "Any other context about this deal"}
        }, "required": ["company"]}
    },
    {
        "name": "get_deal_info",
        "description": "Get full deal details for a specific company — stage, last touchpoint, next action, notes.",
        "input_schema": {"type": "object", "properties": {
            "company": {"type": "string"}
        }, "required": ["company"]}
    },
    {
        "name": "list_deals",
        "description": "List all deals in the pipeline, optionally filtered by stage. Use for 'what's in due diligence?', 'show me the pipeline', 'what deals are we tracking?'.",
        "input_schema": {"type": "object", "properties": {
            "stage": {"type": "string", "description": "Optional filter: sourcing | first_call | due_diligence | passed | invested"}
        }}
    },
    {
        "name": "search_memory",
        "description": "Actively search long-term memory for facts matching a query. Use when you know something should be in memory but it wasn't in the passively-retrieved set, or when Joey explicitly asks to search memories.",
        "input_schema": {"type": "object", "properties": {
            "query": {"type": "string", "description": "What to search for"}
        }, "required": ["query"]}
    },
    {
        "name": "search_deal_database",
        "description": "Search the Africa Big Deal database: 7,000+ historical African startup funding rows (company, round, amount, date, investors). Use for questions about who funded what, when, and for how much. This data is not passively injected — use this tool whenever a question could be about historical funding rounds.",
        "input_schema": {"type": "object", "properties": {
            "query": {"type": "string", "description": "What to search for, e.g. a company name or sector"},
            "max_results": {"type": "integer", "default": 10}
        }, "required": ["query"]}
    },
    {
        "name": "save_procedure",
        "description": "Save a reusable procedure to long-term memory — HOW you handled something, not a fact about WHAT is true. Call this when you work out a non-obvious multi-step approach that will likely recur (e.g. reconciling a Granola note with a mismatched calendar title, a specific sequence for updating a stale deal). Don't call this for one-off or trivial tasks.",
        "input_schema": {"type": "object", "properties": {
            "trigger":   {"type": "string", "description": "Short description of the situation this applies to, e.g. 'Granola note title doesn't match any calendar event'"},
            "procedure": {"type": "string", "description": "The approach that worked, written so future-you can follow it directly"}
        }, "required": ["trigger", "procedure"]}
    },
    {
        "name": "mute_meeting_prep",
        "description": "Stop sending automatic pre-meeting prep briefs for a meeting. Call this in the same turn whenever Joey says a prep brief was unwanted, wrong, or should stop — for a specific meeting or for a category like personal events. This is persistent (survives restarts) — it is the only way to make 'stop prepping X' actually stick.",
        "input_schema": {"type": "object", "properties": {
            "pattern": {"type": "string", "description": "Event title, a distinctive case-insensitive substring of one, or an event ID"},
            "reason":  {"type": "string", "description": "Why this is being muted"}
        }, "required": ["pattern"]}
    },
    {
        "name": "list_meeting_prep_mutes",
        "description": "List all active meeting-prep mutes, with their numbers.",
        "input_schema": {"type": "object", "properties": {}}
    },
    {
        "name": "unmute_meeting_prep",
        "description": "Remove a meeting-prep mute by its number (get the number from list_meeting_prep_mutes first).",
        "input_schema": {"type": "object", "properties": {
            "mute_id": {"type": "integer", "description": "The mute's number from list_meeting_prep_mutes"}
        }, "required": ["mute_id"]}
    }
]

TOOL_FUNCTIONS = {
    "search_gmail_work":    search_gmail_work,
    "read_gmail_work":      read_gmail_work,
    "create_gmail_draft":   create_gmail_draft,
    "list_calendar_events": list_calendar_events,
    "get_calendar_event":   get_calendar_event,
    "create_calendar_event": create_calendar_event,
    "update_calendar_event": update_calendar_event,
    "delete_calendar_event": delete_calendar_event,
    "search_granola":       search_granola,
    "read_granola":         read_granola,
    "search_dropbox":       search_dropbox,
    "read_dropbox_file":    read_dropbox_file,
    "search_drive":         search_drive,
    "read_drive_file":      read_drive_file,
    "update_deal":          update_deal,
    "get_deal_info":        get_deal_info,
    "list_deals":           list_deals,
    "search_memory":        search_memory,
    "search_deal_database": search_deal_database,
    "save_procedure":       save_procedure,
    "mute_meeting_prep":       mute_meeting_prep,
    "list_meeting_prep_mutes": list_meeting_prep_mutes,
    "unmute_meeting_prep":     unmute_meeting_prep,
}

# ── SYSTEM PROMPT ─────────────────────────────────────────────────────────────
BASE_SYSTEM = """You are Dot — a sharp, direct analyst and thinking partner for Joseph (Joey)
Benson-Aruna. He is an investor and advisor working with founders across Francophone
and anglophone Africa, focused on AI, fintech, blockchain, digital financial services,
infrastructure, and technology. He is a partner at DFS Lab (dfs.vc).

Your tools and what they contain:
- search_gmail_work / read_gmail_work: Joey's work Workspace email (founders, investors, DFS Lab)
- create_gmail_draft: create a Gmail draft for Joey to review — NEVER sends automatically
- list_calendar_events / get_calendar_event / create_calendar_event / update_calendar_event / delete_calendar_event: Joey's work Google Calendar (read and write)
- search_granola / read_granola: Call and meeting notes from Granola
- search_dropbox / read_dropbox_file: Dropbox files (pitch decks, PDFs, documents). Reads return 3,000 characters at a time — if a result says more characters remain, call again with the given offset to page further into a long document rather than assuming that's all there is. An image-only PDF (no text layer) is transcribed live by Claude the first time anyone reads it and cached from then on — that first read takes 20-60 seconds, so don't assume the tool has hung.
- search_drive / read_drive_file: Joey's Google Drive (read-only), same 3,000-char paging via offset and the same one-time live transcription for image-only PDFs
- web_search: Real-time web search
- update_deal / get_deal_info / list_deals: deal pipeline (stages: sourcing, first_call, due_diligence, passed, invested)
- search_memory: actively search long-term memory when passive retrieval may have missed something
- search_deal_database: the Africa Big Deal database — 7,000+ historical African startup funding rows (who raised what, when, from whom). Not passively injected; use this tool for any question about historical funding rounds.
- save_procedure: save a reusable HOW-TO when you work out a non-obvious multi-step approach likely to recur — not for facts (those are captured automatically) and not for one-off tasks
- mute_meeting_prep / list_meeting_prep_mutes / unmute_meeting_prep: control automatic pre-meeting prep briefs. If Joey ever says a prep brief was unwanted, wrong, or should stop — for a specific meeting or for a category like personal events — call mute_meeting_prep in that same turn. Do not reply "noted" without calling it; nothing you say is remembered by the prep job unless you write it here.
If Joey asks to be prepped for a meeting or asks for background on a person or company before a call, proactively search Granola, Gmail, Dropbox, and Drive for context and produce a tight prep brief without being asked to use specific tools.

When a question involves a person or company, search Granola first (most recent call context),
then work email. For documents, search Dropbox. For scheduling questions, check calendar.
Search proactively — don't ask permission. For calendar writes (create/update/delete) and email drafts, always confirm the key details in your response before executing — state what you're about to create or change and give Joey a chance to correct it before calling the tool. Email drafts are never sent automatically; Joey reviews and sends from Gmail.
Calendar changes that would notify attendees by email (creating an event with attendees, updating to add attendees, or deleting an event that has attendees) are held automatically by the system and NOT executed until Joey sends /confirm. When a tool result says "PENDING CONFIRMATION", tell Joey exactly what is held and that he must send /confirm to proceed or /cancel to discard. Do not retry the tool call while a confirmation is pending.
A tool failing earlier in this conversation does not mean it is still broken — code changes and restarts happen between messages. Never tell Joey a tool or integration is "still down" or "still broken" based only on an earlier failure in the chat history; always call the tool again in the current turn and report what actually happens now.
Facts retrieved from your long-term memory store appear inside <relevant_memories> tags
at the top of user messages — treat them as background context about Joey and his work.
Reusable procedures from past tasks appear inside <relevant_procedures> tags when relevant —
these are HOW you solved something before, not facts; follow them unless the situation has
clearly changed.
Be direct and specific. No filler, no hedging.
Never narrate your reasoning or process. Do not say 'Let me look that up', 'I'll check', 'Looking at your calendar', 'I'll search for', 'Now I have enough to...', or any similar meta-commentary. Work silently and present results directly.
Keep responses concise and plain. Use bullet points for lists. No emoji headers, no heavy markdown structure, no formatted "report" layout. Write like a sharp colleague — not like an AI generating a document."""

# Frozen system prompt with a cache breakpoint: tools + system cache together
# and stay valid for the whole session. Volatile content (retrieved memories)
# goes into the user turn instead, so it never invalidates this prefix.
SYSTEM_BLOCKS = [{
    "type": "text",
    "text": BASE_SYSTEM,
    "cache_control": {"type": "ephemeral"},
}]

def build_user_content(user_text: str) -> str:
    from datetime import datetime
    from zoneinfo import ZoneInfo
    now = datetime.now(ZoneInfo("America/Toronto"))
    parts = [f"<current_datetime>{now.strftime('%A, %B %-d, %Y %H:%M')} America/Toronto</current_datetime>"]
    memories = retrieve_relevant_memories(user_text, k=15)
    if memories:
        block = "\n".join(f"- {m}" for m in memories)
        parts.append(f"<relevant_memories>\n{block}\n</relevant_memories>")
    procedures = retrieve_relevant_procedures(user_text, k=3)
    if procedures:
        block = "\n".join(f"- {p}" for p in procedures)
        parts.append(f"<relevant_procedures>\n{block}\n</relevant_procedures>")
    parts.append(user_text)
    return "\n\n".join(parts)

# ── CONVERSATION STATE ────────────────────────────────────────────────────────
_BASE           = os.path.dirname(os.path.abspath(__file__))
SESSIONS_DIR    = os.path.join(_BASE, "sessions")
_LEGACY_SESSION = os.path.join(_BASE, "session.json")

def _session_path(name: str) -> str:
    return os.path.join(SESSIONS_DIR, f"{name}.json")

def _migrate_legacy_session():
    """Move session.json → sessions/default.json on first run after upgrade."""
    os.makedirs(SESSIONS_DIR, exist_ok=True)
    if os.path.exists(_LEGACY_SESSION) and not os.path.exists(_session_path("default")):
        shutil.move(_LEGACY_SESSION, _session_path("default"))
        logging.info("Migrated session.json → sessions/default.json")

def repair_history(messages: list) -> list:
    """Drop tool_use blocks whose result was never recorded (process died or
    the API call raised mid-turn) and tool_result blocks orphaned by that —
    either one 400s every subsequent API call. Server-tool results live in the
    same assistant message; client tool results in the following user message.
    Only call between turns: a trailing server_tool_use without a result is
    legitimate while a pause_turn resume is in flight."""
    def ids(content, key):
        return {b.get(key) for b in content if isinstance(b, dict)} if isinstance(content, list) else set()

    while True:
        out, dropped = [], 0
        for i, msg in enumerate(messages):
            content = msg.get("content")
            if not isinstance(content, list):
                out.append(msg)
                continue
            nxt = messages[i + 1].get("content") if i + 1 < len(messages) else None
            prev = out[-1].get("content") if out else None
            answered = ids(content, "tool_use_id") | ids(nxt, "tool_use_id")
            asked = ids(content, "id") | ids(prev, "id")
            kept = [b for b in content if not (isinstance(b, dict) and (
                (b.get("type", "").endswith("tool_use") and b.get("id") not in answered) or
                (b.get("type", "").endswith("tool_result") and b.get("tool_use_id") not in asked)))]
            dropped += len(content) - len(kept)
            if kept:
                out.append(msg if kept == content else {**msg, "content": kept})
        if dropped:
            logging.warning(f"Session repair: dropped {dropped} dangling tool block(s)")
        if out == messages:
            return out
        messages = out

def load_session(name: str = None) -> list:
    path = _session_path(name or _current_session_name)
    try:
        with open(path) as f:
            return repair_history(json.load(f))
    except FileNotFoundError:
        return []
    except Exception as e:
        logging.error(f"Session load error ({name}): {e}")
        return []

def save_session():
    path = _session_path(_current_session_name)
    os.makedirs(SESSIONS_DIR, exist_ok=True)
    try:
        with open(path, "w") as f:
            json.dump(conversation_history, f)
    except Exception as e:
        logging.error(f"Session save error ({_current_session_name}): {e}")

_current_session_name = "default"
_migrate_legacy_session()
conversation_history = load_session()

# ── CLAUDE CALL ───────────────────────────────────────────────────────────────
def _with_cache_breakpoint(messages: list) -> list:
    """Copy of messages with a cache breakpoint on the last content block,
    so the conversation prefix caches turn-over-turn. History itself is never
    mutated — markers would otherwise accumulate past the 4-breakpoint limit."""
    if not messages:
        return messages
    last = json.loads(json.dumps(messages[-1]))
    content = last.get("content")
    if isinstance(content, str):
        content = [{"type": "text", "text": content}]
        last["content"] = content
    if isinstance(content, list) and content and isinstance(content[-1], dict) \
            and content[-1].get("type") in ("text", "tool_result"):
        content[-1]["cache_control"] = {"type": "ephemeral"}
    return messages[:-1] + [last]

def call_claude(container_id=None):
    kwargs = dict(
        model="claude-sonnet-5",
        max_tokens=12000,
        thinking={"type": "adaptive"},
        system=SYSTEM_BLOCKS,
        tools=TOOLS,
        messages=_with_cache_breakpoint(conversation_history),
    )
    if container_id:
        kwargs["container_id"] = container_id
    response = client.messages.create(**kwargs)
    u = response.usage
    logging.info(
        f"usage: in={u.input_tokens} out={u.output_tokens} "
        f"cache_write={getattr(u, 'cache_creation_input_tokens', 0)} "
        f"cache_read={getattr(u, 'cache_read_input_tokens', 0)}"
    )
    return response

# ── AGENT LOOP ────────────────────────────────────────────────────────────────
MAX_TOOL_ITERATIONS = 15

def run_tool(name: str, tool_input: dict):
    """Execute a tool, never raising — a raised exception here would leave a
    dangling tool_use in history and 400 every subsequent API call."""
    global _pending_action
    fn = TOOL_FUNCTIONS.get(name)
    if not fn:
        return f"[{name} is a native tool — handled by API]", False
    if name in _GATED_TOOLS and _needs_confirmation(name, tool_input):
        summary = f"{name}({json.dumps(tool_input, default=str)[:500]})"
        _pending_action = {"name": name, "input": tool_input, "summary": summary}
        return (
            "PENDING CONFIRMATION — this calendar change notifies attendees by email, "
            "so it was NOT executed. Tell Joey exactly what is pending and that he must "
            f"send /confirm to execute it or /cancel to discard it. Pending: {summary}"
        ), False
    try:
        return str(fn(**tool_input)), False
    except Exception as e:
        logging.exception(f"Tool {name} failed")
        return f"Tool error: {e}", True

_whisper_model = None
_prepping_now:  set = set()  # event_ids currently being prepped (in-flight guard, in-memory only)
_prep_notified: set = set()  # event_ids already sent a "couldn't prep" DM this window
_meeting_prep_cal_notified = False  # avoids a DM every 5-min tick during a calendar outage
_INTERNAL_DOMAINS = {'dfslab.net', 'dfs.vc'}
_PREP_TOOL_NAMES  = {'search_gmail_work', 'read_gmail_work', 'search_granola', 'read_granola', 'search_dropbox', 'read_dropbox_file', 'search_drive', 'read_drive_file'}

_pending_action = None  # {"name": str, "input": dict, "summary": str}
_GATED_TOOLS = {"create_calendar_event", "update_calendar_event", "delete_calendar_event"}

def _needs_confirmation(name: str, tool_input: dict) -> bool:
    """True when the call would email third parties via calendar attendee notifications."""
    if name == "create_calendar_event":
        return bool(tool_input.get("attendees"))
    if name == "update_calendar_event":
        return bool(tool_input.get("add_attendees"))
    if name == "delete_calendar_event":
        try:
            e = calendar_work.events().get(
                calendarId='primary', eventId=tool_input.get("event_id", "")
            ).execute()
            return bool(e.get('attendees'))
        except Exception:
            return True  # can't verify — fail safe, require confirmation
    return False

def _get_whisper():
    global _whisper_model
    if _whisper_model is None:
        import whisper
        logging.info("Loading Whisper tiny model...")
        _whisper_model = whisper.load_model("tiny")
        logging.info("Whisper model loaded.")
    return _whisper_model

async def _process_message(update: Update, user_text: str):
    global conversation_history
    # A previous turn that crashed mid-tool-call leaves dangling blocks in the
    # in-memory history without ever hitting disk — repair before building on it.
    conversation_history = repair_history(conversation_history)
    user_content = await asyncio.to_thread(build_user_content, user_text)
    conversation_history.append({"role": "user", "content": user_content})

    if estimate_tokens(conversation_history) > CONTEXT_TOKEN_LIMIT:
        conversation_history = await asyncio.to_thread(compress_history, conversation_history)

    await update.message.chat.send_action("typing")
    try:
        response = await asyncio.to_thread(call_claude)
        container_id = getattr(response, 'container_id', None)

        iterations = 0
        while response.stop_reason in ("tool_use", "pause_turn") and iterations < MAX_TOOL_ITERATIONS:
            iterations += 1
            conversation_history.append({
                "role": "assistant",
                "content": [b.model_dump(exclude_none=True) for b in response.content],
            })
            if response.stop_reason == "tool_use":
                tool_results = []
                for block in response.content:
                    if block.type != "tool_use":
                        continue
                    result, is_error = await asyncio.to_thread(run_tool, block.name, dict(block.input))
                    tr = {"type": "tool_result", "tool_use_id": block.id, "content": result}
                    if is_error:
                        tr["is_error"] = True
                    tool_results.append(tr)
                conversation_history.append({"role": "user", "content": tool_results})
            # pause_turn: re-send so the API can resume from the trailing server_tool_use block.
            await update.message.chat.send_action("typing")
            response = await asyncio.to_thread(call_claude, container_id)
            container_id = getattr(response, 'container_id', None) or container_id

        final = " ".join(b.text for b in response.content if b.type == "text").strip()
        if not final:
            logging.error(f"Empty final response (stop_reason={response.stop_reason})")
            final = ("I hit my output limit before I could answer — the reasoning consumed the "
                     "budget. Ask again, more narrowly." if response.stop_reason == "max_tokens"
                     else "Done.")
        conversation_history.append({"role": "assistant", "content": final})
        save_session()

        for i in range(0, len(final), 4000):
            chunk = final[i:i+4000]
            await send_markdown(update.message.reply_text, chunk)
    except Exception:
        # Roll back to the last clean on-disk save so the next turn doesn't
        # inherit a half-written assistant message and continue it mid-thought.
        conversation_history = load_session()
        raise

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != YOUR_USER_ID:
        return
    user_text = update.message.text
    reply_to = update.message.reply_to_message
    if reply_to and reply_to.text:
        user_text = f"[User is replying to your message: \"{reply_to.text[:500]}\"]\n\n{user_text}"
    await _process_message(update, user_text)

async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != YOUR_USER_ID:
        return

    import tempfile
    voice_file = await context.bot.get_file(update.message.voice.file_id)
    with tempfile.NamedTemporaryFile(suffix='.ogg', delete=False) as tmp:
        await voice_file.download_to_drive(tmp.name)
        tmp_path = tmp.name

    try:
        await update.message.chat.send_action("typing")
        model = await asyncio.to_thread(_get_whisper)
        result = await asyncio.to_thread(model.transcribe, tmp_path)
        text = result["text"].strip()
    except Exception as e:
        logging.exception("Whisper transcription failed")
        await update.message.reply_text(f"Transcription error: {e}")
        return
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    if not text:
        await update.message.reply_text("Couldn't transcribe that — try again.")
        return

    await update.message.reply_text(f"_{text}_", parse_mode="Markdown")
    reply_to = update.message.reply_to_message
    if reply_to and reply_to.text:
        text = f"[User is replying to your message: \"{reply_to.text[:500]}\"]\n\n{text}"
    await _process_message(update, text)

# ── COMMANDS ──────────────────────────────────────────────────────────────────
async def cmd_remember(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    text = " ".join(context.args)
    if text:
        save_memory(text, "manual")
        await update.message.reply_text(f"Saved: {text}")
    else:
        await update.message.reply_text("Usage: /remember [fact]")

async def cmd_memories(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    mems = get_all_memories()[:20]
    if not mems:
        await update.message.reply_text("No memories yet.")
        return
    await update.message.reply_text("\n".join(f"{i+1}. {m}" for i, m in enumerate(mems))[:4000])

async def cmd_forget(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    mems = get_all_memories()[:20]
    try:
        n = int(context.args[0]) - 1
        target = mems[n]
        await asyncio.to_thread(delete_memory, target)
        await update.message.reply_text(f"Deleted: {target}")
    except Exception:
        await update.message.reply_text("Usage: /forget [number from /memories]")

async def cmd_procedures(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    procs = get_all_procedures()[:20]
    if not procs:
        await update.message.reply_text("No procedures saved yet.")
        return
    lines = [f"{i+1}. [{p['trigger']}] {p['procedure']}" for i, p in enumerate(procs)]
    await update.message.reply_text("\n".join(lines)[:4000])

async def cmd_forget_procedure(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    procs = get_all_procedures()[:20]
    try:
        n = int(context.args[0]) - 1
        target = procs[n]
        await asyncio.to_thread(_delete_procedure, target["id"])
        await update.message.reply_text(f"Deleted procedure: [{target['trigger']}] {target['procedure']}")
    except Exception:
        await update.message.reply_text("Usage: /forget_procedure [number from /procedures]")

async def cmd_mutes(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    mutes = list_prep_mutes()[:20]
    if not mutes:
        await update.message.reply_text("No meeting prep mutes set.")
        return
    lines = [
        f"{i+1}. '{m['pattern']}'" + (f" — {m['reason']}" if m['reason'] else "")
        for i, m in enumerate(mutes)
    ]
    await update.message.reply_text("\n".join(lines)[:4000])

async def cmd_unmute(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    mutes = list_prep_mutes()[:20]
    try:
        n = int(context.args[0]) - 1
        target = mutes[n]
        await asyncio.to_thread(delete_prep_mute, target["id"])
        await update.message.reply_text(f"Unmuted: '{target['pattern']}'")
    except Exception:
        await update.message.reply_text("Usage: /unmute [number from /mutes]")

async def cmd_newsession(update, context):
    global conversation_history
    if update.effective_user.id != YOUR_USER_ID: return
    if conversation_history:
        await asyncio.to_thread(extract_and_save_memories, conversation_history)
    conversation_history = []
    save_session()
    await update.message.reply_text("Session cleared. Memories extracted.")

async def cmd_log(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    text = " ".join(context.args)
    if not text:
        await update.message.reply_text("Usage: /log <note or pasted conversation>")
        return
    await update.message.chat.send_action("typing")
    try:
        resp = await asyncio.to_thread(
            client.messages.create,
            model="claude-sonnet-5",
            max_tokens=1500,
            thinking={"type": "disabled"},
            system='Extract facts worth remembering from this note or conversation. Return ONLY a JSON array of strings. Each fact must be self-contained. Focus on people, companies, deals, decisions, commitments. If nothing worth keeping, return [].',
            messages=[{"role": "user", "content": text}]
        )
        facts = parse_json_array(response_text_checked(resp, "cmd_log"))
        saved = 0
        saved_facts = []
        for fact in facts:
            if fact and len(fact) > 10:
                save_memory(fact, "log")
                saved += 1
                saved_facts.append(fact)
        facts_display = "\n".join(f"• {f}" for f in saved_facts)
        reply = f"Logged {saved} fact(s):\n{facts_display}\n\nIf this relates to a deal, just tell me and I'll update the pipeline."
        await update.message.reply_text(reply)
    except Exception as e:
        await update.message.reply_text(f"Log error: {e}")

async def cmd_wrong(update, context):
    """WS-18C — the trial's evidence log. Reply to any Dot message with /wrong [reason]
    to record it. No API call, so this keeps working even when credits are out — exactly
    when Joey most wants to record that something is broken."""
    if update.effective_user.id != YOUR_USER_ID: return
    reply_to = update.message.reply_to_message
    quoted = reply_to.text if reply_to and reply_to.text else ""
    reason = " ".join(context.args)
    try:
        await asyncio.to_thread(save_feedback, quoted, reason)
        await update.message.reply_text("Logged.")
    except Exception as e:
        await update.message.reply_text(f"Feedback log error: {e}")

async def cmd_restart(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    import subprocess
    try:
        result = subprocess.run(
            ["git", "pull", "--ff-only"], cwd=_BASE,
            capture_output=True, text=True, timeout=30,
        )
        pull_output = (result.stdout + result.stderr).strip() or "Already up to date."
    except Exception as e:
        pull_output = f"git pull failed: {e}"
    await update.message.reply_text(f"{pull_output}\n\nRestarting — back in ~10 seconds.")
    import sys
    sys.exit(0)

async def cmd_confirm(update, context):
    global _pending_action
    if update.effective_user.id != YOUR_USER_ID: return
    if not _pending_action:
        await update.message.reply_text("Nothing pending.")
        return
    action, _pending_action = _pending_action, None
    fn = TOOL_FUNCTIONS[action["name"]]
    result = await asyncio.to_thread(lambda: str(fn(**action["input"])))
    await update.message.reply_text(result[:4000])

async def cmd_cancel(update, context):
    global _pending_action
    if update.effective_user.id != YOUR_USER_ID: return
    if _pending_action:
        await update.message.reply_text(f"Cancelled: {_pending_action['summary']}")
        _pending_action = None
    else:
        await update.message.reply_text("Nothing pending.")

async def cmd_switch(update, context):
    global conversation_history, _current_session_name
    if update.effective_user.id != YOUR_USER_ID: return

    name = "_".join(context.args).strip().lower()
    if not name or not re.match(r'^[a-z0-9][a-z0-9_-]*$', name):
        await update.message.reply_text("Usage: /switch <name>  (letters, numbers, hyphens only)\nExample: /switch fundraising")
        return

    if name == _current_session_name:
        await update.message.reply_text(f"Already in '{name}' ({len(conversation_history)} messages).")
        return

    # Save current before switching
    save_session()
    old = _current_session_name
    _current_session_name = name
    conversation_history = load_session(name)

    count = len(conversation_history)
    if count:
        await update.message.reply_text(f"Switched from '{old}' → '{name}' ({count} messages in history).")
    else:
        await update.message.reply_text(f"Switched from '{old}' → '{name}' (new conversation).")

async def cmd_sessions(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    os.makedirs(SESSIONS_DIR, exist_ok=True)
    files = sorted(glob.glob(os.path.join(SESSIONS_DIR, "*.json")))
    if not files:
        await update.message.reply_text(f"Only one conversation: {_current_session_name} (active, {len(conversation_history)} messages)")
        return
    lines = []
    for f in files:
        sname = os.path.splitext(os.path.basename(f))[0]
        try:
            count = len(json.load(open(f)))
        except Exception:
            count = "?"
        marker = " ← active" if sname == _current_session_name else ""
        lines.append(f"• {sname} ({count} msgs){marker}")
    await update.message.reply_text("Conversations:\n" + "\n".join(lines))

async def cmd_search(update, context):
    if update.effective_user.id != YOUR_USER_ID: return
    query = " ".join(context.args)
    if not query:
        await update.message.reply_text("Usage: /search <query>")
        return
    results = await asyncio.to_thread(retrieve_relevant_memories, query, 30)
    if not results:
        await update.message.reply_text("No memories found.")
        return
    text = "\n".join(f"{i+1}. {m}" for i, m in enumerate(results))
    await update.message.reply_text(text[:4000])

# ── ERROR HANDLER ─────────────────────────────────────────────────────────────
async def _notify_owner(context, text: str):
    """Best-effort DM to Joey — must never raise (a failed notify can't crash the job)."""
    try:
        await context.bot.send_message(chat_id=YOUR_USER_ID, text=text[:4000])
    except Exception:
        logging.exception("Failed to notify owner")

# WS-18B (F-37): the credit-balance outage on Aug 12-20 produced one misleading "will
# retry" DM then 911 silent failures over 9 days. This counter is scoped to permanent
# (credit/auth) failures specifically — one DM for the whole outage, not one per event
# per tick — reset on the next success so the next outage gets its own alert.
_consecutive_api_failures = 0

async def _note_api_failure(context, error_msg: str):
    global _consecutive_api_failures
    _consecutive_api_failures += 1
    if _consecutive_api_failures == 1:
        await _notify_owner(
            context,
            f"⚠️ Dot can't reach the Claude API — {error_msg[:300]}. "
            f"Nothing will work until this is resolved. I won't message again about this."
        )

def _note_api_success():
    global _consecutive_api_failures
    _consecutive_api_failures = 0

async def error_handler(update, context):
    logging.error(f"Exception: {context.error}")
    if isinstance(update, Update) and update.effective_message:
        try:
            await update.effective_message.reply_text(f"⚠️ Error: {context.error}")
        except Exception:
            pass
    else:
        await _notify_owner(context, f"⚠️ Background job error: {context.error}")

# ── MEETING PREP BRIEF (runs every 5 min, fires ~30 min before external meetings) ─
def _is_preppable(event: dict, window_start, window_end, internal_domains: set) -> tuple:
    """Pure classification (no API calls, no side effects — WS-19 test target): would
    check_meeting_prep consider this event at all, and why/why not?

    Combines both structural filters from D-10(a): F-29's real "starts within window"
    check (Google Calendar's timeMin/timeMax is an overlap filter, not a starts-within
    filter — an all-day event otherwise matches every 5-minute tick for its whole
    duration) and F-30's external-attendee requirement. Does not know about mutes or
    prep_log — those are stateful and stay in check_meeting_prep."""
    start = event.get('start', {})
    if 'dateTime' not in start:
        return False, "all-day or date-only event"
    from datetime import datetime
    try:
        start_dt = datetime.fromisoformat(start['dateTime'])
    except ValueError:
        return False, "malformed start.dateTime"
    if not (window_start <= start_dt <= window_end):
        return False, "overlaps the query window but does not start within it"
    attendees = event.get('attendees', [])
    external = [
        a for a in attendees
        if not a.get('self', False)
        and not any(a.get('email', '').lower().endswith(f'@{d}') for d in internal_domains)
    ]
    if not external:
        return False, "no external attendees"
    return True, "starts within window with external attendee(s)"

async def check_meeting_prep(context: ContextTypes.DEFAULT_TYPE):
    from datetime import datetime, timezone, timedelta

    now          = datetime.now(timezone.utc)
    window_start = now + timedelta(minutes=25)
    window_end   = now + timedelta(minutes=35)

    global _meeting_prep_cal_notified
    try:
        cal_result = await asyncio.to_thread(
            lambda: calendar_work.events().list(
                calendarId='primary',
                timeMin=window_start.isoformat(),
                timeMax=window_end.isoformat(),
                singleEvents=True,
                orderBy='startTime'
            ).execute()
        )
        events = cal_result.get('items', [])
        _meeting_prep_cal_notified = False
    except Exception as e:
        logging.error(f"Meeting prep calendar error: {e}")
        if not _meeting_prep_cal_notified:
            _meeting_prep_cal_notified = True
            await _notify_owner(context, f"⚠️ Meeting prep couldn't read the calendar: {e}")
        return

    for event in events:
        event_id = event.get('id', '')

        preppable, _reason = _is_preppable(event, window_start, window_end, _INTERNAL_DOMAINS)
        if not preppable:
            continue                      # all-day, malformed, out-of-window, or internal-only (F-29/F-30)

        start_dt = datetime.fromisoformat(event['start']['dateTime'])  # already validated by _is_preppable
        occurrence = start_dt.isoformat()

        title = event.get('summary', 'Untitled')
        if was_prepped(event_id, occurrence) or event_id in _prepping_now:
            continue
        if is_prep_muted(event_id, title):
            mark_prepped(event_id, occurrence, outcome="muted")
            continue

        attendees = event.get('attendees', [])
        external = [
            a for a in attendees
            if not a.get('self', False)
            and not any(a.get('email', '').lower().endswith(f'@{d}') for d in _INTERNAL_DOMAINS)
        ]

        names  = [a.get('displayName') or a.get('email', '') for a in external]
        emails = [a.get('email', '') for a in external]

        prompt = f"""Prepare a meeting brief for Joey (Africa-focused tech investor, DFS Lab).

Meeting: {title} — starting in ~30 minutes
Attendees: {', '.join(names)} ({', '.join(emails)})

Steps:
1. Search Granola for previous call notes with these contacts.
2. Search Gmail for recent email threads with them.
3. Search Dropbox and Drive for any pitch decks, due diligence documents, or investment memos related to these contacts or their company.

Write a tight prep brief:
- Who they are and what their company does (2-3 sentences, from what you find)
- Key context from previous interactions (calls, emails)
- Any relevant documents found (deck highlights, memo conclusions)
- 1-2 suggested talking points or things to follow up on

If this is not a business meeting — a personal appointment, a family event, a travel or
out-of-office block, a focus/hold block, or anything where a prep brief would be unwanted —
respond with exactly SKIP and nothing else. Do not explain. Do not write a brief anyway.

Start immediately with the brief — no preamble, no "I have enough to write...", no "Here's the full picture". Just the content.
Plain prose and bullets. No emoji headers. No markdown section dividers. Write like a sharp colleague gave you a quick verbal rundown."""

        prep_tools = [t for t in TOOLS if t.get('name') in _PREP_TOOL_NAMES]
        msgs = [{"role": "user", "content": prompt}]

        _prepping_now.add(event_id)
        try:
            def _prep_call(cid=None):
                kw = dict(model="claude-sonnet-5", max_tokens=6000,
                          thinking={"type": "adaptive"},
                          tools=prep_tools, messages=msgs)
                if cid:
                    kw["container_id"] = cid
                return client.messages.create(**kw)

            response = await asyncio.to_thread(_prep_call)
            prep_container_id = getattr(response, 'container_id', None)
            iterations = 0
            while response.stop_reason in ("tool_use", "pause_turn") and iterations < 8:
                iterations += 1
                msgs.append({
                    "role": "assistant",
                    "content": [b.model_dump(exclude_none=True) for b in response.content],
                })
                if response.stop_reason == "tool_use":
                    tool_results = []
                    for block in response.content:
                        if block.type == "tool_use":
                            res, is_error = await asyncio.to_thread(run_tool, block.name, dict(block.input))
                            tr = {"type": "tool_result", "tool_use_id": block.id, "content": res}
                            if is_error:
                                tr["is_error"] = True
                            tool_results.append(tr)
                    msgs.append({"role": "user", "content": tool_results})
                response = await asyncio.to_thread(_prep_call, prep_container_id)
                prep_container_id = getattr(response, 'container_id', None) or prep_container_id
            _note_api_success()  # WS-18B: a successful round-trip re-arms the next outage's alert
            text = " ".join(b.text for b in response.content if b.type == "text").strip()
            if text.upper().startswith("SKIP"):
                # Output-side gate (D-10c): the model's own judgment on the Aug 7 evidence was
                # correct and was simply discarded — this is where it now gets acted on. The
                # auto-mute stops a multi-day/long event from burning an API call every tick.
                mark_prepped(event_id, occurrence, outcome="skipped")
                add_prep_mute(event_id, reason=f"auto: model judged '{title}' non-business")
                logging.info(f"Meeting prep: skipped non-business event '{title}'")
                continue
            if text:
                header = f"Prep — {title} (in ~30 min)\n\n"
                chunk = (header + text)[:4000]
                await send_markdown(
                    lambda t, parse_mode=None: context.bot.send_message(
                        chat_id=YOUR_USER_ID, text=t, parse_mode=parse_mode),
                    chunk,
                )
                mark_prepped(event_id, occurrence, outcome="sent")
                _prep_notified.discard(event_id)
        except Exception as e:
            logging.error(f"Meeting prep error for '{title}': {e}")
            msg = str(e)
            if "credit balance" in msg or "authentication" in msg.lower():
                # WS-18B (F-37): permanent/outage-class failure — one DM for the whole
                # outage via the global counter, not one per event per tick (that's what
                # produced 320 failed calls / 0 messages on Aug 15).
                await _note_api_failure(context, msg)
            elif event_id not in _prep_notified:
                _prep_notified.add(event_id)
                await _notify_owner(context, f"⚠️ Couldn't prep '{title}': {msg[:300]} — will retry.")
        finally:
            _prepping_now.discard(event_id)

# ── MAIN ──────────────────────────────────────────────────────────────────────
async def _post_init(app):
    from telegram import BotCommand
    await app.bot.set_my_commands([
        BotCommand("restart",    "Restart the bot"),
        BotCommand("confirm",    "Execute a pending attendee-affecting calendar change"),
        BotCommand("cancel",     "Discard a pending calendar change"),
        BotCommand("switch",     "Switch to a named conversation"),
        BotCommand("sessions",   "List all conversations"),
        BotCommand("log",        "Log a note or WhatsApp forward"),
        BotCommand("remember",   "Save a fact to memory"),
        BotCommand("memories",   "Show recent memories"),
        BotCommand("forget",     "Delete a memory by number"),
        BotCommand("newsession", "Clear current conversation history"),
        BotCommand("search",     "Search long-term memory"),
        BotCommand("mutes",      "List meeting-prep mutes"),
        BotCommand("unmute",     "Remove a meeting-prep mute by number"),
        BotCommand("wrong",      "Reply to a Dot message to log it as wrong"),
    ])

def main():
    app = ApplicationBuilder().token(TELEGRAM_TOKEN).post_init(_post_init).build()
    app.add_handler(CommandHandler("restart",    cmd_restart))
    app.add_handler(CommandHandler("confirm",    cmd_confirm))
    app.add_handler(CommandHandler("cancel",     cmd_cancel))
    app.add_handler(CommandHandler("switch",     cmd_switch))
    app.add_handler(CommandHandler("sessions",   cmd_sessions))
    app.add_handler(CommandHandler("remember",   cmd_remember))
    app.add_handler(CommandHandler("memories",   cmd_memories))
    app.add_handler(CommandHandler("forget",     cmd_forget))
    app.add_handler(CommandHandler("procedures",       cmd_procedures))
    app.add_handler(CommandHandler("forget_procedure", cmd_forget_procedure))
    app.add_handler(CommandHandler("mutes",      cmd_mutes))
    app.add_handler(CommandHandler("unmute",     cmd_unmute))
    app.add_handler(CommandHandler("newsession", cmd_newsession))
    app.add_handler(CommandHandler("log",        cmd_log))
    app.add_handler(CommandHandler("wrong",      cmd_wrong))
    app.add_handler(CommandHandler("search",     cmd_search))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    app.add_handler(MessageHandler(filters.VOICE, handle_voice))
    app.add_error_handler(error_handler)
    app.job_queue.run_repeating(check_meeting_prep, interval=300, first=60)
    print("Dot is running — Granola, Dropbox, Gmail, Calendar (work)...")
    app.run_polling()

if __name__ == "__main__":
    main()

